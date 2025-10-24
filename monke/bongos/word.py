"""Word bongo implementation.

Creates, updates, and deletes test Word documents via the Microsoft Graph API.
"""

import asyncio
import io
import time
import uuid
from typing import Any, Dict, List, Optional

import httpx
from monke.bongos.base_bongo import BaseBongo
from monke.generation.word import generate_documents_content
from monke.utils.logging import get_logger

# Try to import python-docx for Word document creation
try:
    from docx import Document
    from docx.shared import Pt

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

GRAPH = "https://graph.microsoft.com/v1.0"


class WordBongo(BaseBongo):
    """Bongo for Word that creates test entities for E2E testing.

    Key responsibilities:
    - Create test Word documents in OneDrive
    - Add content with verification tokens
    - Update documents to test incremental sync
    - Delete documents to test deletion detection
    - Clean up all test data
    """

    connector_type = "word"

    def __init__(self, credentials: Dict[str, Any], **kwargs):
        super().__init__(credentials)
        self.access_token: str = credentials["access_token"]
        self.entity_count: int = int(kwargs.get("entity_count", 3))
        self.openai_model: str = kwargs.get("openai_model", "gpt-4.1-mini")
        self.rate_limit_delay = float(kwargs.get("rate_limit_delay_ms", 500)) / 1000.0
        self.logger = get_logger("word_bongo")

        # Track created resources for cleanup
        self._test_documents: List[Dict[str, Any]] = []
        self._last_req = 0.0

        if not HAS_PYTHON_DOCX:
            raise ImportError(
                "python-docx is required for Word bongo. Install with: pip install python-docx"
            )

    async def create_entities(self) -> List[Dict[str, Any]]:
        """Create test Word documents in OneDrive."""
        self.logger.info(f"🥁 Creating {self.entity_count} Word test documents")
        out: List[Dict[str, Any]] = []

        # Generate tokens for each document
        tokens = [uuid.uuid4().hex[:8] for _ in range(self.entity_count)]

        # Generate document content
        test_name = f"TestDoc_{uuid.uuid4().hex[:8]}"
        filenames, document_content = await generate_documents_content(
            self.openai_model, tokens, test_name
        )

        self.logger.info(f"📄 Generated {len(document_content)} documents")

        async with httpx.AsyncClient(base_url=GRAPH, timeout=60) as client:
            for i, (filename, doc_content, token) in enumerate(
                zip(filenames, document_content, tokens)
            ):
                await self._pace()

                # Sanitize filename for OneDrive (remove illegal characters)
                safe_filename = self._sanitize_filename(filename)

                # Create Word document file
                doc_bytes = self._create_word_file(doc_content)

                # Upload to OneDrive
                self.logger.info(f"📤 Uploading Word document: {safe_filename}")
                upload_url = f"/me/drive/root:/{safe_filename}:/content"

                r = await client.put(
                    upload_url,
                    headers={
                        "Authorization": f"Bearer {self.access_token}",
                        "Content-Type": (
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    },
                    content=doc_bytes,
                )

                if r.status_code not in (200, 201):
                    self.logger.error(f"Upload failed {r.status_code}: {r.text}")
                    r.raise_for_status()

                doc_file = r.json()
                doc_id = doc_file["id"]

                self.logger.info(f"✅ Uploaded document: {doc_id} - {filename}")

                ent = {
                    "type": "document",
                    "id": doc_id,
                    "filename": safe_filename,
                    "title": doc_content.title,
                    "token": token,
                    "expected_content": token,
                }
                out.append(ent)
                self._test_documents.append(ent)
                self.created_entities.append({"id": doc_id, "name": safe_filename})

                self.logger.info(f"📝 Document '{safe_filename}' created with token: {token}")

        self.logger.info(f"✅ Created {len(self._test_documents)} Word documents")
        return out

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename by removing illegal characters for OneDrive.

        Args:
            filename: Original filename

        Returns:
            Sanitized filename safe for OneDrive
        """
        # Replace illegal characters: \ / : * ? " < > |
        illegal_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
        safe_name = filename
        for char in illegal_chars:
            safe_name = safe_name.replace(char, '_')

        # Remove leading/trailing spaces and dots
        safe_name = safe_name.strip('. ')

        # Limit length to 200 characters (OneDrive has a 400 char limit for full path)
        if len(safe_name) > 200:
            # Keep the extension
            name, ext = safe_name.rsplit('.', 1) if '.' in safe_name else (safe_name, '')
            safe_name = name[:195] + '.' + ext if ext else name[:200]

        return safe_name

    def _create_word_file(self, doc_content: Any) -> bytes:
        """Create a Word document file with the given content.

        Args:
            doc_content: WordDocumentContent object

        Returns:
            Bytes of the Word document
        """
        doc = Document()

        # Add title
        title = doc.add_heading(doc_content.title, level=0)
        title.alignment = 1  # Center alignment

        # Add summary section
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(doc_content.summary)

        # Add main content
        doc.add_heading("Content", level=1)

        # Split content by paragraphs and add them
        paragraphs = doc_content.content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                # Set font size for readability
                for run in p.runs:
                    run.font.size = Pt(11)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    async def update_entities(self) -> List[Dict[str, Any]]:
        """Update Word documents by appending new content with same tokens."""
        if not self._test_documents:
            return []

        self.logger.info(f"🥁 Updating {min(2, len(self._test_documents))} Word documents")
        updated = []

        async with httpx.AsyncClient(base_url=GRAPH, timeout=60) as client:
            for ent in self._test_documents[: min(2, len(self._test_documents))]:
                await self._pace()

                # Download current document
                download_url = f"/me/drive/items/{ent['id']}/content"
                r = await client.get(download_url, headers=self._hdrs())

                if r.status_code != 200:
                    self.logger.warning(
                        f"Failed to download document: {r.status_code} - {r.text[:200]}"
                    )
                    continue

                # Load existing document
                doc = Document(io.BytesIO(r.content))

                # Append update section with token
                doc.add_heading("Update Section", level=1)
                doc.add_paragraph(
                    f"This document has been updated. Update token: {ent['token']}\n"
                    f"Updated content to verify incremental sync functionality."
                )

                # Save updated document
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                updated_bytes = buffer.getvalue()

                # Upload updated document
                upload_url = f"/me/drive/items/{ent['id']}/content"
                r = await client.put(
                    upload_url,
                    headers={
                        "Authorization": f"Bearer {self.access_token}",
                        "Content-Type": (
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    },
                    content=updated_bytes,
                )

                if r.status_code in (200, 201):
                    updated.append({**ent, "updated": True})
                    self.logger.info(
                        f"📝 Updated document '{ent['filename']}' with token: {ent['token']}"
                    )
                else:
                    self.logger.warning(
                        f"Failed to update document: {r.status_code} - {r.text[:200]}"
                    )

                # Brief delay between updates
                await asyncio.sleep(0.5)

        return updated

    async def delete_entities(self) -> List[str]:
        """Delete all test documents."""
        return await self.delete_specific_entities(self._test_documents)

    async def delete_specific_entities(self, entities: List[Dict[str, Any]]) -> List[str]:
        """Delete specific Word documents with retry for locked files."""
        if not entities:
            # If no specific entities provided, delete all tracked documents
            entities = self._test_documents

        if not entities:
            return []

        self.logger.info(f"🥁 Deleting {len(entities)} Word documents")
        deleted: List[str] = []

        async with httpx.AsyncClient(base_url=GRAPH, timeout=30) as client:
            for ent in entities:
                try:
                    await self._pace()

                    # Try to delete the document with retry for locked files (423)
                    max_retries = 3
                    retry_delay = 2.0  # seconds

                    for attempt in range(max_retries):
                        r = await client.delete(
                            f"/me/drive/items/{ent['id']}", headers=self._hdrs()
                        )

                        if r.status_code == 204:
                            deleted.append(ent["id"])
                            self.logger.info(f"✅ Deleted document: {ent.get('filename', ent['id'])}")

                            # Remove from tracking
                            if ent in self._test_documents:
                                self._test_documents.remove(ent)
                            break  # Success, exit retry loop

                        elif r.status_code == 423 and attempt < max_retries - 1:
                            # Resource is locked, wait and retry
                            self.logger.warning(
                                f"⏳ Document locked (423), retrying in {retry_delay}s "
                                f"(attempt {attempt + 1}/{max_retries}): {ent.get('filename', ent['id'])}"
                            )
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2  # Exponential backoff
                        else:
                            # Other error or max retries reached
                            self.logger.warning(
                                f"Delete failed: {r.status_code} - {r.text[:200]}"
                            )
                            break  # Exit retry loop on non-retryable error

                except Exception as e:
                    self.logger.warning(f"Delete error for {ent.get('filename', ent['id'])}: {e}")

        return deleted

    async def cleanup(self):
        """Comprehensive cleanup of all test resources."""
        self.logger.info("🧹 Starting comprehensive Word cleanup")

        cleanup_stats = {
            "documents_deleted": 0,
            "errors": 0,
        }

        try:
            async with httpx.AsyncClient(base_url=GRAPH, timeout=30) as client:
                # Delete tracked test documents
                if self._test_documents:
                    self.logger.info(f"🗑️  Deleting {len(self._test_documents)} tracked documents")
                    deleted = await self.delete_specific_entities(self._test_documents.copy())
                    cleanup_stats["documents_deleted"] += len(deleted)

                # Search for and cleanup any orphaned test documents
                await self._cleanup_orphaned_documents(client, cleanup_stats)

            self.logger.info(
                f"🧹 Cleanup completed: {cleanup_stats['documents_deleted']} "
                f"documents deleted, {cleanup_stats['errors']} errors"
            )
        except Exception as e:
            self.logger.error(f"❌ Error during comprehensive cleanup: {e}")

    async def _cleanup_orphaned_documents(
        self, client: httpx.AsyncClient, stats: Dict[str, Any]
    ):
        """Find and delete orphaned test documents from previous runs."""
        try:
            await self._pace()
            r = await client.get("/me/drive/root/children", headers=self._hdrs())

            if r.status_code == 200:
                files = r.json().get("value", [])

                # Find test Word documents
                test_documents = [
                    f
                    for f in files
                    if f.get("name", "").startswith("Monke_")
                    and f.get("name", "").endswith(".docx")
                ]

                if test_documents:
                    self.logger.info(
                        f"🔍 Found {len(test_documents)} orphaned test documents"
                    )
                    for doc in test_documents:
                        try:
                            await self._pace()
                            del_r = await client.delete(
                                f"/me/drive/items/{doc['id']}",
                                headers=self._hdrs(),
                            )
                            if del_r.status_code == 204:
                                stats["documents_deleted"] += 1
                                self.logger.info(
                                    f"✅ Deleted orphaned document: {doc.get('name', 'Unknown')}"
                                )
                            else:
                                stats["errors"] += 1
                        except Exception as e:
                            stats["errors"] += 1
                            self.logger.warning(
                                f"⚠️  Failed to delete document {doc['id']}: {e}"
                            )
        except Exception as e:
            self.logger.warning(f"⚠️  Could not search for orphaned documents: {e}")

    def _hdrs(self) -> Dict[str, str]:
        """Get standard headers for Graph API requests."""
        return {
            "Authorization": f"Bearer {self.access_token}",
            "Content-Type": "application/json",
        }

    async def _pace(self):
        """Rate limiting helper."""
        now = time.time()
        if (delta := now - self._last_req) < self.rate_limit_delay:
            await asyncio.sleep(self.rate_limit_delay - delta)
        self._last_req = time.time()

